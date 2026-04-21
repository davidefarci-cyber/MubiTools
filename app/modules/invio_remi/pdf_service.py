"""Servizio generazione PDF da template DOCX per invio REMI."""

import asyncio
import logging
import os
import shutil
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Mm
from docx.oxml.ns import qn

from app.modules.invio_remi import settings_service

logger = logging.getLogger(__name__)


def _replace_tag_in_paragraph(paragraph, tag: str, replacement: str) -> bool:
    """Sostituisce un tag nel testo di un paragrafo, gestendo run frammentati.

    Returns True se la sostituzione è avvenuta.
    """
    full_text = paragraph.text
    if tag not in full_text:
        return False

    new_text = full_text.replace(tag, replacement)

    # Cancella tutti i run tranne il primo, poi imposta il testo nel primo
    for i in range(len(paragraph.runs) - 1, 0, -1):
        paragraph.runs[i].clear()
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.text = new_text

    return True


def _replace_tag_with_table(paragraph, tag: str, remi_codes: list[str]) -> bool:
    """Sostituisce il tag <REMI> con una tabella di codici REMI.

    Returns True se la sostituzione è avvenuta.
    """
    if tag not in paragraph.text:
        return False

    # Rimuovi il testo del paragrafo che contiene il tag
    for run in paragraph.runs:
        run.clear()
    if not paragraph.runs:
        paragraph.text = ""

    # Inserisci la tabella subito dopo il paragrafo corrente
    parent = paragraph._element.getparent()
    doc_for_table = paragraph._element

    # Crea tabella OPC
    tbl = _build_remi_table(paragraph, remi_codes)
    doc_for_table.addnext(tbl)

    return True


def _build_remi_table(paragraph, remi_codes: list[str]):
    """Costruisce un elemento tabella OPC con una colonna e una riga per codice REMI."""
    from docx.oxml import OxmlElement

    tbl = OxmlElement("w:tbl")

    # Proprietà tabella
    tbl_pr = OxmlElement("w:tblPr")
    tbl_style = OxmlElement("w:tblStyle")
    tbl_style.set(qn("w:val"), "TableGrid")
    tbl_pr.append(tbl_style)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")
    tbl_pr.append(tbl_w)

    # Allineamento tabella
    tbl_jc = OxmlElement("w:jc")
    tbl_jc.set(qn("w:val"), "left")
    tbl_pr.append(tbl_jc)

    tbl.append(tbl_pr)

    # Griglia tabella
    tbl_grid = OxmlElement("w:tblGrid")
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), "5000")
    tbl_grid.append(grid_col)
    tbl.append(tbl_grid)

    # Stile bordi
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "999999")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # Righe
    for code in remi_codes:
        tr = OxmlElement("w:tr")

        tc = OxmlElement("w:tc")

        # Padding cella
        tc_pr = OxmlElement("w:tcPr")
        tc_mar = OxmlElement("w:tcMar")
        for side in ("top", "left", "bottom", "right"):
            mar = OxmlElement(f"w:{side}")
            mar.set(qn("w:w"), "80")
            mar.set(qn("w:type"), "dxa")
            tc_mar.append(mar)
        tc_pr.append(tc_mar)
        tc.append(tc_pr)

        # Paragrafo dentro la cella
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")

        # Font coerente
        r_pr = OxmlElement("w:rPr")
        r_sz = OxmlElement("w:sz")
        r_sz.set(qn("w:val"), "20")  # 10pt
        r_pr.append(r_sz)
        r_sz_cs = OxmlElement("w:szCs")
        r_sz_cs.set(qn("w:val"), "20")
        r_pr.append(r_sz_cs)
        r.append(r_pr)

        t = OxmlElement("w:t")
        t.text = code
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
        tbl.append(tr)

    return tbl


def _replace_tags_in_document(doc: Document, replacements: dict[str, str], remi_codes: list[str]) -> None:
    """Sostituisce tutti i tag nel documento (paragrafi e tabelle)."""

    # Processa paragrafi del body
    for paragraph in list(doc.paragraphs):
        if "<REMI>" in paragraph.text:
            _replace_tag_with_table(paragraph, "<REMI>", remi_codes)
        else:
            for tag, value in replacements.items():
                _replace_tag_in_paragraph(paragraph, tag, value)

    # Processa tabelle esistenti nel documento
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "<REMI>" in paragraph.text:
                        _replace_tag_with_table(paragraph, "<REMI>", remi_codes)
                    else:
                        for tag, value in replacements.items():
                            _replace_tag_in_paragraph(paragraph, tag, value)

    # Processa header e footer
    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            if header and header.is_linked_to_previous is False:
                for paragraph in header.paragraphs:
                    for tag, value in replacements.items():
                        _replace_tag_in_paragraph(paragraph, tag, value)
        for footer in (section.footer, section.first_page_footer):
            if footer and footer.is_linked_to_previous is False:
                for paragraph in footer.paragraphs:
                    for tag, value in replacements.items():
                        _replace_tag_in_paragraph(paragraph, tag, value)


def format_date_for_display(date_str: str) -> str:
    """Converte data da YYYY-MM-DD a DD/MM/YYYY."""
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return dt.strftime("%d/%m/%Y")
    except (ValueError, TypeError):
        return date_str


async def generate_pdf(
    company_name: str,
    pec_address: str,
    effective_date: str,
    remi_codes: list[str],
) -> bytes:
    """Genera un PDF dal template DOCX tramite conversione con soffice.

    Args:
        company_name: Ragione sociale del distributore.
        pec_address: Indirizzo PEC del distributore.
        effective_date: Data decorrenza (formato YYYY-MM-DD).
        remi_codes: Lista di codici REMI da inserire.

    Returns:
        Contenuto binario del PDF generato.

    Raises:
        RuntimeError: Se soffice non è disponibile o la conversione fallisce.
        FileNotFoundError: Se il template DOCX non esiste.
    """
    template_path = settings_service.get_template_path()
    if not template_path.exists():
        raise FileNotFoundError("Template DOCX non trovato")

    # Verifica disponibilità soffice
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        raise RuntimeError(
            "soffice non trovato nel PATH di sistema. "
            "Installare OpenOffice/LibreOffice con il pacchetto writer "
            "(es. apt install libreoffice-writer)"
        )

    # Crea directory temporanea per il lavoro
    tmp_dir = tempfile.mkdtemp(prefix="remi_pdf_")
    try:
        # Apri e modifica il template
        doc = Document(str(template_path))

        replacements = {
            "<NOME_DL>": company_name,
            "<PEC_DL>": pec_address,
            "<DATA_DECORRENZA>": format_date_for_display(effective_date),
            "<DATA>": datetime.now().strftime("%d/%m/%Y"),
        }

        _replace_tags_in_document(doc, replacements, remi_codes)

        # Salva il DOCX modificato
        tmp_docx = os.path.join(tmp_dir, "document.docx")
        doc.save(tmp_docx)

        # Converti in PDF con soffice headless
        logger.info("Avvio conversione PDF: %s -> %s", tmp_docx, tmp_dir)
        process = await asyncio.create_subprocess_exec(
            soffice_path,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", tmp_dir,
            tmp_docx,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await process.communicate()

        stdout_txt = stdout.decode("utf-8", errors="replace").strip()
        stderr_txt = stderr.decode("utf-8", errors="replace").strip()

        if stdout_txt:
            logger.info("soffice stdout: %s", stdout_txt)
        if stderr_txt:
            logger.warning("soffice stderr: %s", stderr_txt)

        if process.returncode != 0:
            logger.error(
                "soffice conversione fallita (exit code %d): stdout=%s stderr=%s",
                process.returncode, stdout_txt, stderr_txt,
            )
            raise RuntimeError(
                f"Conversione PDF fallita (exit code {process.returncode}): {stderr_txt or stdout_txt}"
            )

        # Leggi il PDF generato
        tmp_pdf = os.path.join(tmp_dir, "document.pdf")
        if not os.path.exists(tmp_pdf):
            logger.error(
                "PDF non generato nonostante exit code 0. stdout=%s stderr=%s",
                stdout_txt, stderr_txt,
            )
            raise RuntimeError(
                f"PDF non generato da soffice. "
                f"Verificare che il pacchetto writer sia installato. "
                f"Output soffice: {stdout_txt or stderr_txt}"
            )

        with open(tmp_pdf, "rb") as f:
            pdf_bytes = f.read()

        logger.info("PDF generato con successo (%d bytes)", len(pdf_bytes))
        return pdf_bytes

    finally:
        # Pulizia file temporanei
        shutil.rmtree(tmp_dir, ignore_errors=True)
